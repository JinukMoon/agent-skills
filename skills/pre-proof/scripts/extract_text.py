#!/usr/bin/env python
"""Extract text from a PDF or Word file for pre-proof analysis.

Goal: produce a faithful, location-tagged plain-text dump that PRESERVES every
character exactly as it appears in the source -- including unicode subscripts,
smart quotes, ligatures, non-breaking spaces, etc. Those "ugly" characters are
precisely the anomalies the pre-proof is hunting for, so we must NOT normalize
or clean anything here.

Output: a sibling <stem>_extracted.txt with page/paragraph markers so findings
can be located. PDF pages are marked "===== PAGE n ====="; Word paragraphs are
not page-bound, so they are marked "[P0001]" etc. (paragraph index), and table
cells as "[T<table>.<row>.<col>]".

Usage:
    python extract_text.py <input.pdf|input.docx> [output.txt]
"""
import sys
import os


def extract_pdf(path):
    import fitz  # PyMuPDF
    doc = fitz.open(path)
    out = []
    for i, page in enumerate(doc, 1):
        out.append(f"\n===== PAGE {i} =====")
        # "text" preserves reading order and unicode codepoints as-is.
        out.append(page.get_text("text"))
    doc.close()
    return "\n".join(out)


def extract_docx(path):
    import docx
    d = docx.Document(path)
    out = []
    for i, para in enumerate(d.paragraphs, 1):
        txt = para.text
        if txt.strip() == "":
            continue
        out.append(f"[P{i:04d}] {txt}")
    # tables (captions, data) are easy to miss otherwise
    for ti, table in enumerate(d.tables, 1):
        out.append(f"\n----- TABLE {ti} -----")
        for ri, row in enumerate(table.rows, 1):
            for ci, cell in enumerate(row.cells, 1):
                ctxt = cell.text.strip()
                if ctxt:
                    out.append(f"[T{ti}.{ri}.{ci}] {ctxt}")
    return "\n".join(out)


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)
    path = sys.argv[1]
    if not os.path.isfile(path):
        print(f"ERROR: file not found: {path}")
        sys.exit(1)
    ext = os.path.splitext(path)[1].lower()
    if ext == ".pdf":
        text = extract_pdf(path)
    elif ext in (".docx",):
        text = extract_docx(path)
    elif ext == ".doc":
        print("ERROR: legacy .doc not supported. Convert to .docx or PDF first "
              "(e.g. libreoffice --headless --convert-to docx file.doc).")
        sys.exit(1)
    else:
        print(f"ERROR: unsupported extension {ext}. Use .pdf or .docx")
        sys.exit(1)

    if len(sys.argv) >= 3:
        outpath = sys.argv[2]
    else:
        stem = os.path.splitext(path)[0]
        outpath = stem + "_extracted.txt"

    with open(outpath, "w", encoding="utf-8") as f:
        f.write(text)

    nlines = text.count("\n") + 1
    print(f"OK: extracted {len(text)} chars / {nlines} lines -> {outpath}")


if __name__ == "__main__":
    main()
